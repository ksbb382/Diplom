"""
Build ФИНАЛЬНЫЙ_ДИПЛОМ_ВОРОБЬЕВ.docx by:
1. Take ВКР Рыжкин (template) -> keep titul + задание + аннотация рус/англ + СОДЕРЖАНИЕ table.
   Cut everything from his ВВЕДЕНИЕ onwards.
2. Replace Рыжкин's data with Воробьев's data (name, theme, content, annotation).
3. Replace ОГЛАВЛЕНИЕ table with Воробьев's structure (manual, like Рыжкин — page numbers as "00").
4. Append ВКР Воробьев.docx (Введение → Заключение) via docxcompose.
5. Add inline references [N] in 12-15 places throughout the text.
6. Append "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" with 25 sources in Рыжкин's style.
"""

import os
import re
import shutil
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxcompose.composer import Composer

REPO = "/projects/sandbox/Diplom"
RYZHKIN = os.path.join(REPO, "ВКР Рыжкин А. А. 237.docx")
VOROBYEV = os.path.join(REPO, "ВКР Воробьев.docx")
OUT = os.path.join(REPO, "diplom", "ФИНАЛЬНЫЙ_ДИПЛОМ_ВОРОБЬЕВ.docx")

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# ====================================================================
# STEP 1: Build template from Рыжкин (titul + задание + аннотация + СОДЕРЖАНИЕ)
# ====================================================================

def build_template():
    """Open Рыжкин and crop to keep only the front matter (before ВВЕДЕНИЕ).
    Replace text with Воробьев's data and rebuild the СОДЕРЖАНИЕ table."""
    work_path = "/tmp/voro_template.docx"
    shutil.copy(RYZHKIN, work_path)
    
    doc = Document(work_path)
    body = doc.element.body
    
    # Find the cutoff point: paragraph "ВВЕДЕНИЕ" with style ds-markdown-paragraph
    # at position 201 in original. Actually, we keep up to and including the СОДЕРЖАНИЕ table 
    # (around par 167 + table). Then cut everything after.
    
    # Find "ВВЕДЕНИЕ" paragraph - it's the first one with this exact text after the ToC table
    cutoff_elem = None
    seen_soderzhanie = False
    
    for child in list(body.iterchildren()):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
            if text == 'СОДЕРЖАНИЕ':
                seen_soderzhanie = True
                continue
            if seen_soderzhanie and text == 'ВВЕДЕНИЕ':
                cutoff_elem = child
                break
        # also catch the first "Глава 1." after ToC (Рыжкин uses this for the body)
    
    if cutoff_elem is None:
        # Fallback: look for "ВВЕДЕНИЕ" as Heading 1
        for child in list(body.iterchildren()):
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
                if text == 'ВВЕДЕНИЕ':
                    cutoff_elem = child
                    break
    
    if cutoff_elem is None:
        raise RuntimeError("Could not find cutoff point (ВВЕДЕНИЕ)")
    
    # Remove cutoff_elem and everything after it (until sectPr)
    sectPr = body.find(qn('w:sectPr'))
    
    # Collect elements to remove (cutoff_elem and onwards until sectPr)
    to_remove = []
    found_cutoff = False
    for child in list(body.iterchildren()):
        if child is cutoff_elem:
            found_cutoff = True
        if found_cutoff:
            tag = child.tag.split('}')[-1]
            if tag == 'sectPr':
                continue
            to_remove.append(child)
    
    for el in to_remove:
        body.remove(el)
    
    print(f"[Template] Removed {len(to_remove)} elements after СОДЕРЖАНИЕ")
    
    # Now replace text in titul + задание + аннотация
    replace_voroboyev_data(doc)
    
    # Rebuild СОДЕРЖАНИЕ table
    rebuild_toc_table(doc)
    
    # Add page break after СОДЕРЖАНИЕ to ensure ВВЕДЕНИЕ starts on new page
    add_page_break_at_end(body, sectPr)
    
    doc.save(work_path)
    print(f"[Template] Saved to {work_path}")
    return work_path


def replace_in_runs(para, replacements):
    """Replace text in a paragraph's runs using a dict {old: new}.
    Tries to preserve formatting by editing first run that contains the old text."""
    full_text = ''.join(t.text or '' for t in para._element.iter(qn('w:t')))
    
    for old, new in replacements.items():
        if old not in full_text:
            continue
        new_full = full_text.replace(old, new)
        full_text = new_full
        
        # Write back: put new_full into first <w:t>, clear the rest in this paragraph
        t_nodes = list(para._element.iter(qn('w:t')))
        if t_nodes:
            t_nodes[0].text = new_full
            for tn in t_nodes[1:]:
                tn.text = ""


def replace_voroboyev_data(doc):
    """Walk through all paragraphs and replace Рыжкин's data with Воробьев's."""
    
    # Map: replace these strings everywhere
    replacements = {
        'Рыжкину Артемию Алексеевичу': 'Воробьеву Павлу Алексеевичу',
        'Рыжкин Артемий Алексеевич': 'Воробьев Павел Алексеевич',
        '(Рыжкин А. А.)': '(Воробьев П. А.)',
        '(Рыжкин А.А.)': '(Воробьев П.А.)',
        'Рыжкин А. А.': 'Воробьев П. А.',
        'Рыжкин А.А.': 'Воробьев П.А.',
        'Рыжкин': 'Воробьев',
        'Инвестиционный проект малого предприятия': 'Инвестиционный проект производственного предприятия по выпуску комбинированных пожарных извещателей',
    }
    
    for para in doc.paragraphs:
        replace_in_runs(para, replacements)
    
    # Also walk through tables
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, replacements)
    
    # ----- Replace the "5 Содержание ВКР" block (paragraphs 71-86 in Рыжкин) -----
    # Find Глава 1 entry and replace with Воробьев's structure
    new_content = [
        ('Введение', 'Основной текст (14)'),
        ('Глава 1. Технико-экономическое обоснование проекта', 'ds-markdown-paragraph'),
        ('1.1 Обоснование целей открытия предприятия', 'ds-markdown-paragraph'),
        ('1.2 Анализ рынка', 'ds-markdown-paragraph'),
        ('1.3 Выбор инструмента анализа деятельности предприятия', 'ds-markdown-paragraph'),
        ('Глава 2. Постановка задачи', 'ds-markdown-paragraph'),
        ('Глава 3. Проектная часть', 'ds-markdown-paragraph'),
        ('3.1 Исходные данные', 'ds-markdown-paragraph'),
        ('3.2 Календарный план работ', 'ds-markdown-paragraph'),
        ('3.3 Управление качеством', 'ds-markdown-paragraph'),
        ('3.4 Внесение данных в Project Expert', 'ds-markdown-paragraph'),
        ('3.5 Анализ чувствительности', 'ds-markdown-paragraph'),
        ('Глава 4. Выводы по эффективности предложения', 'ds-markdown-paragraph'),
        ('Заключение', 'Основной текст (14)'),
        ('Список использованных источников', 'Основной текст (14)'),
    ]
    
    body = doc.element.body
    
    # Find paragraph "Введение" (after "5 Содержание ВКР") in remaining doc and replace block
    paragraphs = list(body.iter(qn('w:p')))
    
    intro_idx = None
    for i, p_el in enumerate(paragraphs):
        text = ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()
        if text == 'Введение':
            # Check that previous para is "5 Содержание ВКР"
            if i > 0:
                prev_text = ''.join(t.text or '' for t in paragraphs[i-1].iter(qn('w:t'))).strip()
                if 'Содержание ВКР' in prev_text:
                    intro_idx = i
                    break
    
    if intro_idx is None:
        print("[Template] Warning: Could not find '5 Содержание ВКР' / 'Введение' block")
        return
    
    # Find end of content list — paragraph "Приложения" or empty + "6 Перечень..."
    end_idx = None
    for j in range(intro_idx, len(paragraphs)):
        text = ''.join(t.text or '' for t in paragraphs[j].iter(qn('w:t'))).strip()
        if 'Приложения' == text or text.startswith('6 Перечень'):
            end_idx = j
            break
    
    if end_idx is None:
        end_idx = intro_idx + 16  # fallback
    
    print(f"[Template] Replacing content list pars {intro_idx}..{end_idx-1} with new structure")
    
    # Take template paragraph (paragraphs[intro_idx]) for style reference; clone it for each new entry
    template_para = paragraphs[intro_idx]
    parent = template_para.getparent()
    
    # Insert new paragraphs BEFORE template_para
    for text, style_name in new_content:
        new_p = OxmlElement('w:p')
        # pPr with style
        pPr = OxmlElement('w:pPr')
        pStyle = OxmlElement('w:pStyle')
        # Resolve style ID from name
        style_id = None
        for s in doc.styles:
            if s.name == style_name:
                style_id = s.style_id
                break
        if style_id:
            pStyle.set(qn('w:val'), style_id)
            pPr.append(pStyle)
        new_p.append(pPr)
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = text
        r.append(t)
        new_p.append(r)
        template_para.addprevious(new_p)
    
    # Now remove old block (intro_idx ... end_idx-1) — these are still in `paragraphs` list
    for old_p in paragraphs[intro_idx:end_idx]:
        if old_p.getparent() is not None:
            old_p.getparent().remove(old_p)
    
    # ----- Replace "6 Перечень графического (демонстрационного) материала" -----
    # Find this paragraph and replace the 8 items below it
    paragraphs = list(body.iter(qn('w:p')))
    perech_idx = None
    for i, p_el in enumerate(paragraphs):
        text = ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()
        if text.startswith('6 Перечень'):
            perech_idx = i
            break
    
    if perech_idx is not None:
        # Find next non-empty paragraph indicating end (usually empty para then "Дата выдачи задания")
        end_idx2 = None
        for j in range(perech_idx + 1, len(paragraphs)):
            text = ''.join(t.text or '' for t in paragraphs[j].iter(qn('w:t'))).strip()
            if text.startswith('Дата выдачи') or text.startswith('Задание принял'):
                end_idx2 = j
                break
        
        if end_idx2 is None:
            end_idx2 = perech_idx + 12
        
        new_demo = [
            '1. Структурно-логическая схема инвестиционного проекта',
            '2. Календарный план реализации проекта (диаграмма Ганта)',
            '3. Блок-схема технологического процесса сборки и контроля качества',
            '4. Динамика объёма рынка пожарной автоматики РФ',
            '5. Структура капитальных вложений (CAPEX)',
            '6. Динамика доходов и расходов проекта',
            '7. График окупаемости проекта',
            '8. Анализ чувствительности по NPV, DPP, PI',
        ]
        
        # Insert new BEFORE first old item, then remove old items
        if perech_idx + 1 < len(paragraphs):
            anchor = paragraphs[perech_idx + 1]
            for line in new_demo:
                new_p = OxmlElement('w:p')
                pPr = OxmlElement('w:pPr')
                # Use same style as original first item (Normal usually)
                anchor_pStyle = anchor.find(qn('w:pPr') + '/' + qn('w:pStyle'))
                if anchor_pStyle is not None:
                    pStyle = OxmlElement('w:pStyle')
                    pStyle.set(qn('w:val'), anchor_pStyle.get(qn('w:val')))
                    pPr.append(pStyle)
                new_p.append(pPr)
                r = OxmlElement('w:r')
                t = OxmlElement('w:t')
                t.text = line
                r.append(t)
                new_p.append(r)
                anchor.addprevious(new_p)
            
            # Remove old items
            for old_p in paragraphs[perech_idx + 1:end_idx2]:
                if old_p.getparent() is not None:
                    old_p.getparent().remove(old_p)
        
        print(f"[Template] Replaced graphical materials list")
    
    # ----- Replace АННОТАЦИЯ paragraphs (rus + eng) -----
    replace_annotation(doc)


def replace_annotation(doc):
    """Replace Russian and English annotation."""
    
    annotation_ru = [
        ("Объектом исследования в выпускной квалификационной работе выступает малое производственное предприятие "
         "ООО «СигналПром», ориентированное на серийный выпуск комбинированных дымо-тепловых пожарных извещателей. "
         "Предмет исследования — процесс экономического обоснования инвестиционного проекта производства извещателя "
         "ИПК 212/101-1-СП «Сигнал» с применением модели гибридной автоматизации производства."),
        
        ("Цель работы — разработка инвестиционного проекта малого производственного предприятия по выпуску "
         "комбинированных пожарных извещателей с проведением комплексной оценки экономической эффективности проекта "
         "на основе расчёта интегральных показателей и анализа чувствительности."),
        
        ("Методология исследования включает расчётно-аналитический метод, методы сравнительного анализа, элементы "
         "сценарного моделирования, графический метод и приёмы бизнес-планирования. В качестве основного "
         "инструмента финансового моделирования применяется специализированный программный комплекс "
         "Project Expert 6 Holding. Информационная база включает нормативные правовые акты Российской Федерации в "
         "области пожарной безопасности, статистические материалы Росстата и Банка России, а также аналитические "
         "материалы Минпромторга России и Корпорации МСП."),
        
        ("Результатом работы является разработанный инвестиционный проект производственного предприятия "
         "со следующими параметрами: бюджет капитальных вложений (CAPEX) — 4 974 482 руб., чистый дисконтированный "
         "доход (NPV) — 2 461 669 руб., внутренняя норма доходности (IRR) — 35,67%, индекс прибыльности (PI) — 1,31, "
         "дисконтированный срок окупаемости (DPB) — 30 месяцев. Анализ чувствительности подтвердил устойчивость "
         "проекта к снижению объёма сбыта в пределах 15–20%."),
        
        ("Практическая значимость исследования заключается в возможности использования разработанной модели "
         "в качестве основы для подготовки бизнес-плана, обоснования заявки на привлечение финансирования "
         "(в том числе через программы льготного лизинга и грантов «Мой бизнес»), а также для последующей "
         "адаптации проекта к конкретным региональным и рыночным условиям."),
    ]
    
    annotation_en = [
        ("The object of the research is a small manufacturing enterprise OOO «SignalProm» specialised in the "
         "production of combined smoke-and-heat fire detectors. The subject of the research is the process of "
         "economic justification of an investment project for the production of the IPK 212/101-1-SP «Signal» "
         "fire detector using a hybrid automation model."),
        
        ("The purpose of the thesis is to develop an investment project for a small manufacturing enterprise "
         "producing combined fire detectors and to perform a comprehensive evaluation of its economic efficiency "
         "based on integral performance indicators and sensitivity analysis."),
        
        ("The research methodology is based on analytical calculations, comparative analysis, scenario "
         "modelling, graphical methods and business-planning tools. The main financial modelling instrument is "
         "the specialised software complex Project Expert 6 Holding. The information base includes Russian "
         "regulatory and legal acts on fire safety, statistical data of Rosstat and the Bank of Russia, and "
         "analytical materials of the Ministry of Industry and Trade and the SME Corporation."),
        
        ("The thesis results in a developed investment project with the following key parameters: capital "
         "expenditures (CAPEX) — 4,974,482 rubles, net present value (NPV) — 2,461,669 rubles, internal rate of "
         "return (IRR) — 35.67%, profitability index (PI) — 1.31, discounted payback period (DPB) — 30 months. "
         "The sensitivity analysis has confirmed the resilience of the project to a 15–20% reduction in sales volume."),
        
        ("The practical value of the work lies in the possibility of using the proposed model as a basis for "
         "a business plan, for financing applications (including through preferential leasing programmes and "
         "«My Business» grants), and for adapting the project to specific regional and market conditions."),
    ]
    
    body = doc.element.body
    paragraphs = list(body.iter(qn('w:p')))
    
    # Find АННОТАЦИЯ paragraph
    anno_ru_idx = None
    anno_en_idx = None
    for i, p_el in enumerate(paragraphs):
        text = ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()
        if text == 'АННОТАЦИЯ':
            anno_ru_idx = i
        elif text == 'ANNOTATION':
            anno_en_idx = i
    
    if anno_ru_idx is None or anno_en_idx is None:
        print("[Template] Warning: Could not find АННОТАЦИЯ/ANNOTATION")
        return
    
    # Replace Russian annotation paragraphs (5 paragraphs after АННОТАЦИЯ)
    _replace_annotation_block(paragraphs, anno_ru_idx, annotation_ru, 5)
    
    # Re-find paragraphs (indices may have shifted; we only changed text in place though)
    # Replace English annotation
    paragraphs = list(body.iter(qn('w:p')))
    anno_en_idx = None
    for i, p_el in enumerate(paragraphs):
        text = ''.join(t.text or '' for t in p_el.iter(qn('w:t'))).strip()
        if text == 'ANNOTATION':
            anno_en_idx = i
            break
    
    if anno_en_idx is not None:
        _replace_annotation_block(paragraphs, anno_en_idx, annotation_en, 5)
    
    print("[Template] Replaced АННОТАЦИЯ rus + ANNOTATION eng")


def _replace_annotation_block(paragraphs, header_idx, new_paragraphs, expected_count):
    """Replace `expected_count` paragraphs after the header with new_paragraphs.
    Uses run-text replacement to keep formatting."""
    
    # Strategy: find next `expected_count` non-empty paragraphs after header_idx
    target_pars = []
    j = header_idx + 1
    while j < len(paragraphs) and len(target_pars) < expected_count:
        text = ''.join(t.text or '' for t in paragraphs[j].iter(qn('w:t'))).strip()
        if text:
            target_pars.append(paragraphs[j])
        j += 1
    
    # Replace text in each target paragraph
    for target_p, new_text in zip(target_pars, new_paragraphs):
        t_nodes = list(target_p.iter(qn('w:t')))
        if t_nodes:
            t_nodes[0].text = new_text
            for tn in t_nodes[1:]:
                tn.text = ""


def rebuild_toc_table(doc):
    """Replace the static СОДЕРЖАНИЕ table (Table 0) with Воробьев's structure.
    Pages stay as '00' — like Рыжкин's. User fills them in Word manually."""
    
    if not doc.tables:
        print("[Template] Warning: No tables, cannot rebuild ToC")
        return
    
    tbl = doc.tables[0]
    
    # Voroboyev's contents: (text, page)
    voro_toc = [
        ('Введение……………………………………………………………………………', '00'),
        ('Глава 1. Технико-экономическое обоснование проекта…………………...', '00'),
        ('1.1 Обоснование целей открытия предприятия…………………………...', '00'),
        ('1.2 Анализ рынка……………………………………………………………...', '00'),
        ('1.3 Выбор инструмента анализа деятельности предприятия……………', '00'),
        ('Глава 2. Постановка задачи…………………………………………………...', '00'),
        ('Глава 3. Проектная часть………………………………………………………', '00'),
        ('3.1 Исходные данные…………………………………………………………', '00'),
        ('3.2 Календарный план работ………………………………………………...', '00'),
        ('3.3 Управление качеством……………………………………………………', '00'),
        ('3.4 Внесение данных в Project Expert……………………………………...', '00'),
        ('3.5 Анализ чувствительности………………………………………………..', '00'),
        ('Глава 4. Выводы по эффективности предложения…………………………', '00'),
        ('Заключение……………………………………………………………………...', '00'),
        ('Список использованных источников……………………………………….', '00'),
    ]
    
    # Original table has 21 rows; we need 15. Adjust rows.
    # Easier: replace text in existing rows (up to min(15, current_rows)),
    # then remove extra rows.
    
    n_target = len(voro_toc)
    n_current = len(tbl.rows)
    
    # Replace text in existing rows
    for i, (text, page) in enumerate(voro_toc[:min(n_target, n_current)]):
        row = tbl.rows[i]
        # Cell 0: text, Cell 1: page
        for col_idx, content in enumerate([text, page]):
            if col_idx >= len(row.cells):
                break
            cell = row.cells[col_idx]
            # Get all paragraphs in cell, put text in first, clear rest
            cell_pars = cell.paragraphs
            if cell_pars:
                # Replace text in first paragraph, preserving its style
                first_para = cell_pars[0]
                t_nodes = list(first_para._element.iter(qn('w:t')))
                if t_nodes:
                    t_nodes[0].text = content
                    for tn in t_nodes[1:]:
                        tn.text = ""
                else:
                    # No <w:t> nodes — create a run
                    r = OxmlElement('w:r')
                    t = OxmlElement('w:t')
                    t.text = content
                    r.append(t)
                    first_para._element.append(r)
                # Clear extra paragraphs in cell
                for extra_p in cell_pars[1:]:
                    for tn in extra_p._element.iter(qn('w:t')):
                        tn.text = ""
    
    # If current rows > target, remove excess rows
    if n_current > n_target:
        rows_to_remove = list(tbl.rows[n_target:])
        for row in rows_to_remove:
            tr = row._tr
            tr.getparent().remove(tr)
        print(f"[Template] Removed {len(rows_to_remove)} extra ToC rows")
    
    # If current < target — would need to add rows (not our case here, n_target=15 < n_current=21)
    
    print(f"[Template] Rebuilt ToC table: {n_target} rows")


def add_page_break_at_end(body, sectPr):
    """Add a page break after the last content paragraph (before sectPr)."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    if sectPr is not None:
        sectPr.addprevious(p)
    else:
        body.append(p)


# ====================================================================
# STEP 2: Add inline references [N] in Воробьев's text
# ====================================================================

def add_inline_references(work_path):
    """Replace existing [1], [2], [8], [9], [10] markers with proper [1]..[5],
    then add additional inline references in chapters 1.2, 2, 3.1, 3.3, 3.4."""
    
    doc = Document(work_path)
    
    # Step A: Renumber existing markers ([8]→[3], [9]→[4], [10]→[5])
    # so we get sequential [1]..[5] for the Introduction.
    rename_map = {'[8]': '[3]', '[9]': '[4]', '[10]': '[5]'}
    
    for para in doc.paragraphs:
        full = ''.join(t.text or '' for t in para._element.iter(qn('w:t')))
        if not any(k in full for k in rename_map):
            continue
        new_full = full
        for old, new in rename_map.items():
            new_full = new_full.replace(old, new)
        if new_full != full:
            t_nodes = list(para._element.iter(qn('w:t')))
            if t_nodes:
                t_nodes[0].text = new_full
                for tn in t_nodes[1:]:
                    tn.text = ""
    
    # Step B: Add inline references in specific places
    # Each entry: (marker_text_to_find, ref_to_append)
    # Find a unique substring in the text, append [N] before its sentence end.
    inline_refs = [
        # Ch 1.2 Анализ рынка
        ('доля отечественных производителей в государственном, муниципальном и крупном корпоративном секторах приблизилась', '[6]'),
        ('Изменение №1 к СП 484, вступившее в силу с 1 сентября 2025 года', '[7]'),
        ('обновлённый национальный стандарт ГОСТ Р 53325-2024', '[8]'),
        # Ch 2 Постановка задачи
        ('Методические рекомендации по оценке эффективности инвестиционных проектов', '[9]'),
        ('Федерального закона «Об инвестиционной деятельности', '[10]'),
        # Ch 3.1 Приборы и оборудование
        ('сертификат соответствия ТР ЕАЭС 043/2017', '[11]'),
        ('ГОСТ IEC 61340-5-1-2019', '[12]'),
        # Ch 3.3 Управление качеством
        ('IPC-A-610', '[13]'),
        ('ГОСТ Р ИСО 2859', '[14]'),
        # Ch 3.4 Project Expert
        ('Project Expert 6 Holding', '[15]'),
        # Ch 1.1 - tech сбор / Минпромторг
        ('технологического сбора на импортируемую', '[16]'),
        ('Корпорации МСП', '[17]'),
    ]
    
    inserted = set()  # track which refs we've actually inserted
    for para in doc.paragraphs:
        t_nodes = list(para._element.iter(qn('w:t')))
        if not t_nodes:
            continue
        full = ''.join(t.text or '' for t in t_nodes)
        
        for marker, ref in inline_refs:
            if ref in inserted:
                continue
            if marker not in full:
                continue
            # Append ref at the end of the sentence containing marker
            # Find marker position, then find next sentence terminator
            m_pos = full.find(marker)
            # Find end of sentence after marker (period, but skip dots in abbreviations like "т. д." — use period followed by space/uppercase or end)
            search_pos = m_pos + len(marker)
            # Look for next ". " or ".\n" or end
            sentence_end = -1
            i = search_pos
            while i < len(full):
                if full[i] == '.':
                    # Check next char
                    if i + 1 >= len(full) or full[i+1] in ' \n\t':
                        sentence_end = i
                        break
                i += 1
            
            if sentence_end == -1:
                # Just put at end of paragraph
                sentence_end = len(full) - 1 if full.endswith('.') else len(full)
            
            # Insert ref BEFORE the period
            if sentence_end < len(full) and full[sentence_end] == '.':
                new_full = full[:sentence_end] + ' ' + ref + full[sentence_end:]
            else:
                new_full = full[:sentence_end] + ' ' + ref + full[sentence_end:]
            
            t_nodes[0].text = new_full
            for tn in t_nodes[1:]:
                tn.text = ""
            
            full = new_full
            t_nodes = list(para._element.iter(qn('w:t')))
            inserted.add(ref)
            print(f"[Refs] Inserted {ref} after '{marker[:50]}...'")
    
    print(f"[Refs] Total inline refs inserted: {len(inserted)} (refs 6-17)")
    print(f"[Refs] Refs not placed: {sorted(set(r for _, r in inline_refs) - inserted)}")
    
    doc.save(work_path)


# ====================================================================
# STEP 3: Compose: template + Воробьев
# ====================================================================

def compose_thesis(template_path, voro_path):
    """Compose template + Воробьев's content, save to OUT."""
    
    # Make a working copy of Воробьев to modify (add refs)
    voro_work = "/tmp/voro_body.docx"
    shutil.copy(voro_path, voro_work)
    add_inline_references(voro_work)
    
    base = Document(template_path)
    composer = Composer(base)
    
    voro_doc = Document(voro_work)
    composer.append(voro_doc)
    
    composer.save(OUT)
    print(f"[Compose] Saved {OUT}")


# ====================================================================
# STEP 4: Add bibliography (25 sources) at the end
# ====================================================================

BIBLIOGRAPHY = [
    # ---- Federal Laws (нормативно-правовые) ----
    "Российская Федерация. Законы. Об инвестиционной деятельности в Российской Федерации, осуществляемой в форме капитальных вложений: Федеральный закон № 39-ФЗ: [принят Государственной думой 15 июля 1998 года: одобрен Советом Федерации 17 июля 1998 года: редакция от 28 декабря 2024 года]. – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_22142/ (дата обращения: 14.05.2026).",
    
    "Российская Федерация. Законы. О развитии малого и среднего предпринимательства в Российской Федерации: Федеральный закон № 209-ФЗ: [принят Государственной думой 6 июля 2007 года: одобрен Советом Федерации 11 июля 2007 года: редакция от 8 марта 2026 года]. – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_52144/ (дата обращения: 14.05.2026).",
    
    "Российская Федерация. Законы. Технический регламент о требованиях пожарной безопасности: Федеральный закон № 123-ФЗ: [принят Государственной думой 4 июля 2008 года: одобрен Советом Федерации 11 июля 2008 года: действующая редакция]. – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_78699/ (дата обращения: 14.05.2026).",
    
    "Российская Федерация. Законы. Налоговый кодекс Российской Федерации (часть вторая): Федеральный закон № 117-ФЗ: [принят Государственной думой 19 июля 2000 года: одобрен Советом Федерации 26 июля 2000 года: действующая редакция]. Глава 26.2 «Упрощённая система налогообложения». – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_28165/ (дата обращения: 14.05.2026).",
    
    "Российская Федерация. Законы. О промышленной политике в Российской Федерации: Федеральный закон № 488-ФЗ: [принят Государственной думой 16 декабря 2014 года: одобрен Советом Федерации 25 декабря 2014 года: редакция от 28 декабря 2024 года]. – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_173119/ (дата обращения: 14.05.2026).",
    
    # ---- Технические регламенты, ГОСТы, СП ----
    "Технический регламент Евразийского экономического союза «О требованиях к средствам обеспечения пожарной безопасности и пожаротушения»: ТР ЕАЭС 043/2017: [утверждён Решением Совета Евразийской экономической комиссии от 23 июня 2017 года № 40]. – Текст: электронный. – URL: https://www.eaeunion.org/ (дата обращения: 14.05.2026).",
    
    "ГОСТ Р 53325-2024. Технические средства пожарной автоматики. Общие технические требования. Методы испытаний. – Москва: Российский институт стандартизации, 2024. – Текст: электронный. – URL: https://www.gostinfo.ru/ (дата обращения: 14.05.2026).",
    
    "СП 484.1311500.2020. Системы противопожарной защиты. Системы пожарной сигнализации и автоматизация систем противопожарной защиты. Нормы и правила проектирования: [с Изменением № 1 от 1 сентября 2025 года]. – Москва: МЧС России, 2020. – Текст: электронный. – URL: https://www.mchs.gov.ru/dokumenty/svody-pravil (дата обращения: 14.05.2026).",
    
    "ГОСТ IEC 61340-5-1-2019. Электростатика. Часть 5-1. Защита электронных устройств от электростатических явлений. Общие требования. – Москва: Стандартинформ, 2019. – Текст: электронный.",
    
    "IPC-A-610. Acceptability of Electronic Assemblies: Class 2. – Bannockburn, IL: IPC – Association Connecting Electronics Industries, 2020. – Text: print.",
    
    "ГОСТ Р ИСО 2859-1-2007. Статистические методы. Процедуры выборочного контроля по альтернативному признаку. – Москва: Стандартинформ, 2008. – Текст: электронный.",
    
    # ---- Подзаконные акты ----
    "Об утверждении правил предоставления субсидий российским организациям на компенсацию части затрат на производство и реализацию пилотных партий средств производства и материалов: Постановление Правительства Российской Федерации от 17 февраля 2016 года № 109 (с изменениями от 2025 года). – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный.",
    
    "Об утверждении перечня товаров, в отношении которых допускается параллельный импорт: Приказ Минпромторга России от 8 октября 2024 года № 4611. – Текст: электронный. – URL: https://minpromtorg.gov.ru/ (дата обращения: 14.05.2026).",
    
    "Методические рекомендации по оценке эффективности инвестиционных проектов: утв. Минэкономики РФ, Минфином РФ, Госстроем РФ 21.06.1999 № ВК 477. – Доступ из справочно-правовой системы «КонсультантПлюс». – Текст: электронный. – URL: https://www.consultant.ru/document/cons_doc_LAW_28224/ (дата обращения: 14.05.2026).",
    
    # ---- Книги/учебники ----
    "Виленский, П. Л. Оценка эффективности инвестиционных проектов: теория и практика: учебно-практическое пособие / П. Л. Виленский, В. Н. Лившиц, С. А. Смоляк. – 5-е изд., перераб. и доп. – Москва: Дело, 2015. – 1104 с. – Текст: непосредственный.",
    
    "Лимитовский, М. А. Инвестиционные проекты и реальные опционы на развивающихся рынках: учебно-практическое пособие / М. А. Лимитовский. – 5-е изд., перераб. и доп. – Москва: Юрайт, 2019. – 486 с. – Текст: непосредственный.",
    
    # ---- Электронные ресурсы (статистика) ----
    "Банк России. Ключевая ставка Банка России: динамика // Официальный сайт Центрального банка Российской Федерации. – Текст: электронный. – URL: https://www.cbr.ru/hd_base/keyrate/ (дата обращения: 14.05.2026).",
    
    "Малое и среднее предпринимательство в России. 2024: статистический сборник / Федеральная служба государственной статистики (Росстат). – Москва: Росстат, 2024. – Текст: электронный. – URL: https://rosstat.gov.ru/storage/mediabank/Mal_pred_2024.pdf (дата обращения: 14.05.2026).",
    
    "ФНС России. Единый реестр субъектов малого и среднего предпринимательства: официальный сервис // Официальный сайт Федеральной налоговой службы. – Текст: электронный. – URL: https://www.nalog.gov.ru/rn77/related_activities/regbusiness/ (дата обращения: 14.05.2026).",
    
    "Корпорация МСП. Льготный лизинг оборудования для малого и среднего бизнеса: программа поддержки // Официальный сайт АО «Корпорация «МСП». – Текст: электронный. – URL: https://corpmsp.ru/to-business/lizingovaya-podderzhka/ (дата обращения: 14.05.2026).",
    
    "Минпромторг России. Стимулирование производства электронной продукции: национальный проект // Официальный сайт Министерства промышленности и торговли Российской Федерации. – Текст: электронный. – URL: https://minpromtorg.gov.ru/ (дата обращения: 14.05.2026).",
    
    "МЧС России. Государственная статистика по пожарам и их последствиям // Официальный сайт МЧС России. – Текст: электронный. – URL: https://www.mchs.gov.ru/dokumenty/statistika (дата обращения: 14.05.2026).",
    
    # ---- ПО и поставщики ЭКБ ----
    "Project Expert: программный комплекс для разработки бизнес-планов и анализа инвестиционных проектов: руководство пользователя / ООО «Эксперт Системс». – Москва: Эксперт Системс, 2024. – Текст: электронный. – URL: https://www.expert-systems.ru/ (дата обращения: 14.05.2026).",
    
    "ООО «Резонит»: каталог услуг контрактного производства печатных плат и SMT-монтажа // Официальный сайт компании «Резонит». – Текст: электронный. – URL: https://www.rezonit.ru/ (дата обращения: 14.05.2026).",
    
    "АО «ЧИП и ДИП»: каталог электронных компонентов и измерительного оборудования // Официальный сайт компании «ЧИП и ДИП». – Текст: электронный. – URL: https://www.chipdip.ru/ (дата обращения: 14.05.2026).",
]


def add_bibliography():
    """Open the composed file and append the bibliography section."""
    doc = Document(OUT)
    
    body = doc.element.body
    sectPr = body.find(qn('w:sectPr'))
    
    # Add page break before bibliography
    p_break = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p_break.append(r)
    if sectPr is not None:
        sectPr.addprevious(p_break)
    else:
        body.append(p_break)
    
    # Add Heading 1: СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ
    heading_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    # Find Heading 1 style
    h1_id = None
    for s in doc.styles:
        if s.name == 'Heading 1':
            h1_id = s.style_id
            break
    if h1_id:
        pStyle.set(qn('w:val'), h1_id)
        pPr.append(pStyle)
    heading_p.append(pPr)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ'
    r.append(t)
    heading_p.append(r)
    if sectPr is not None:
        sectPr.addprevious(heading_p)
    else:
        body.append(heading_p)
    
    # Add 25 numbered sources as Normal-style paragraphs
    # Find a "Normal-like" style; use exact paragraph properties similar to Рыжкин body
    for idx, src in enumerate(BIBLIOGRAPHY, start=1):
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        # Indent first line + justify (like Рыжкин)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:firstLine'), '709')  # ~1.25 cm
        pPr.append(ind)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'both')
        pPr.append(jc)
        # Spacing: 1.5
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '360')  # 360 twentieths = 1.5 line spacing
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        p.append(pPr)
        
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # Font: Liberation Serif 14 pt (28 half-points)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Liberation Serif')
        rFonts.set(qn('w:hAnsi'), 'Liberation Serif')
        rFonts.set(qn('w:cs'), 'Liberation Serif')
        rPr.append(rFonts)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '28')
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), '28')
        rPr.append(szCs)
        r.append(rPr)
        
        t = OxmlElement('w:t')
        t.text = f"{idx}. {src}"
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        
        if sectPr is not None:
            sectPr.addprevious(p)
        else:
            body.append(p)
    
    print(f"[Bibliography] Added СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ ({len(BIBLIOGRAPHY)} sources)")
    
    doc.save(OUT)


# ====================================================================
# MAIN
# ====================================================================

def main():
    template_path = build_template()
    compose_thesis(template_path, VOROBYEV)
    add_bibliography()
    print(f"\n[DONE] {OUT}")
    print(f"Size: {os.path.getsize(OUT) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
